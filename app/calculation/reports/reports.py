import logging
import os
import csv
import io
import docx
import pandas as pd
import numpy as np
import inspect

import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.gridspec as gridspec
# from matplotlib import font_manager as fm, rcParams
from typing import IO
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from datetime import datetime
from docx.shared import Inches

from app.infrastructure.database.models.substance import FlammableMaterialModel

log = logging.getLogger(__name__)


def get_csv_file(data, name_file) -> csv:
    file = str(name_file)
    try:
        with open(file, 'w', newline='') as file_w:
            csv_writer = csv.writer(
                file_w, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
            csv_writer.writerows(data)

    except csv.Error:
        print(f"Ошибка в CSV-файле на строке {reader.line_num}: {csv.Error}")
    except UnicodeDecodeError:
        print("Ошибка декодирования. Возможно, файл имеет другую кодировку.")
    except FileNotFoundError:
        print("Файл не найден.")
    except IOError:
        print(f"Ошибка ввода-вывода при работе с файлом. {IOError}")
    except Exception as e:
        print(f"Непредвиденная ошибка: {e}")


def get_csv_bt_file(data) -> bytes:
    output = io.StringIO()
    with output as file_w:
        writer = csv.writer(file_w, dialect='excel', delimiter=';',
                            quotechar='"', quoting=csv.QUOTE_MINIMAL)
        writer.writerows(data)
        byte_file = output.getvalue().encode("ANSI")
    return byte_file


# def get_xlsx_bt_file(data) -> bytes:

#     writer = pd.ExcelWriter('simple-report.xlsx', engine='xlsxwriter')
#     df.to_excel(writer, index=False)
#     df_footer.to_excel(writer, startrow=6, index=False)
#     writer.save()
#     output = io.StringIO()
#     with output as file_w:
#         writer = csv.writer(file_w, dialect='excel', delimiter=';',
#                             quotechar='"', quoting=csv.QUOTE_MINIMAL)
#         writer.writerows(data)
#         byte_file = output.getvalue().encode("ANSI")
#     return byte_file


def get_data_fire_load(file_path: str, subname: str, data: FlammableMaterialModel):
    # Открываем шаблон документа
    document = docx.Document(file_path)
    paragraphs = document.paragraphs
    table_1 = document.tables[0]
    # header = paragraphs[0]
    paragraphs[0].runs[-1].text = data.substance_name
    paragraphs[1].runs[-1].text = data.description if data.description != None else data.substance_name
    paragraphs[6].runs[-1].text = '\n'.join(data.data_source.split('\n')[:-1])
    # paragraphs[1].runs[-1].text = kwargs[subname].get("description", "-")
    # paragraphs[6].runs[-1].text = data[subname].get("data_source", "-")

    # Заполняем таблицу с параметрами горючей нагрузки

    md_list = [data.lower_heat_of_combustion,
               data.linear_flame_velocity,
               data.specific_burnout_rate,
               data.combustion_efficiency,
               data.smoke_forming_ability,
               data.oxygen_consumption,
               data.carbon_dioxide_output,
               data.carbon_monoxide_output,
               data.hydrogen_chloride_output,
               data.critical_heat_flux]

    for string in range(1, len(table_1.rows)):
        table_1.rows[string].cells[2].text = str(md_list[string - 1])
        # data[subname].get(model_data.get(string), "-")
        table_1.rows[string].cells[3].text = "[1]"
    document.save(
        rf"temp_files/temp_data/flammable_load_{subname.lower()}.docx")


def get_report_analytics_model(name: str, file: str | IO[bytes]):

    document = docx.Document()
    # paragraphs = document.paragraphs
    # table_1 = document.tables[0]
    # # header = paragraphs[0]
    # # paragraphs[0].runs[-1].text = data.substance_name
    # # paragraphs[1].runs[-1].text = data.description if data.description != None else data.substance_name
    # # paragraphs[6].runs[-1].text = '\n'.join(data.data_source.split('\n')[:-1])
    # # paragraphs[1].runs[-1].text = kwargs[subname].get("description", "-")
    # # paragraphs[6].runs[-1].text = data[subname].get("data_source", "-")

    # # Заполняем таблицу с параметрами горючей нагрузки

    # md_list = [data.lower_heat_of_combustion,
    #            data.linear_flame_velocity,
    #            data.specific_burnout_rate,
    #            data.combustion_efficiency,
    #            data.smoke_forming_ability,
    #            data.oxygen_consumption,
    #            data.carbon_dioxide_output,
    #            data.carbon_monoxide_output,
    #            data.hydrogen_chloride_output,
    #            data.critical_heat_flux]

    # for string in range(1, len(table_1.rows)):
    #     table_1.rows[string].cells[2].text = str(md_list[string - 1])
    #     # data[subname].get(model_data.get(string), "-")
    #     table_1.rows[string].cells[3].text = "[1]"
    # document.save(
    #     rf"temp_files/temp_data/flammable_load_{subname.lower()}.docx")
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet')
    document.add_paragraph(
        'first item in ordered list', style='List Number')

    document.add_picture(
        image_path_or_stream=file, width=Inches(5))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'))

    table = document.add_table(rows=1, cols=3, style='Table Normal')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()
    document.add_page_break()

    document.save(
        rf"temp_files/temp_data/report_analytics_model_{name.lower()}.docx")
