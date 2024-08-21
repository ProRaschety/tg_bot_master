import logging
import os
import datetime
import re
import docx

from scipy.interpolate import RectBivariateSpline, interp1d
from CoolProp import CoolProp

log = logging.getLogger(__name__)


def clean_name(name):
    """
    Преобразовать имя строки в буквенно-цифровой строчный регистр

    """
    parsed = re.sub(r'\W+', '', name.lower())
    return parsed


def get_distance_at_value(x_values, y_values, value):
    func_value = interp1d(y_values, x_values, kind='linear',
                          bounds_error=False, fill_value=0)
    return func_value(value)


def get_value_at_distance(x_values, y_values, distance):
    func_distance = interp1d(x_values, y_values, kind='linear',
                             bounds_error=False, fill_value=0)
    return func_distance(distance)


def update_number_picture(doc_path: str,
                          figure_first_count: int = 1,
                          table_first_count: int = 1):
    """Функция для сквозной нумерации рисунков в документе."""
    # Открываем документ
    doc = docx.Document(doc_path)
    figure_count = figure_first_count  # Счетчик для рисунков
    table_count = table_first_count  # Счетчик для таблиц
    num_table = 0
    num_para = 0
    amount_pictures = []

    # Проходим по всем элементам тела документа
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Если это параграф
            para = doc.paragraphs[num_para]
            # para = doc.paragraphs[doc.element.body.index(element)]
            # Ищем подписи к рисункам
            matches = re.findall(r'Рисунок \d+\.|Рис\. \d+\.', para.text)
            for match in matches:
                # print('Номер рисунка №', figure_count)
                # Заменяем старую подпись на новую
                new_text = re.sub(r'Рисунок \d+\.|Рис\. \d+\.',
                                  f'Рисунок {figure_count}.', match)
                para.text = para.text.replace(match, new_text)
                # Очищаем параграф и добавляем новый текст
                # para.clear()  # Очищаем параграф
                # para.add_run(para.text.replace(match, new_text))  # Добавляем новый текст
                figure_count += 1
            num_para += 1

        elif element.tag.endswith('tbl'):  # Если это таблица
            # Проходим по всем таблицам
            # table = doc.tables[doc.element.body.index(element)]
            table = doc.tables[num_table]
            num_table += 1
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(
                        r'Рисунок \d+\.|Рис\. \d+\.', cell.text)
                    for match in matches:
                        # print('Номер рисунка №', figure_count)
                        new_text = re.sub(
                            r'Рисунок \d+\.|Рис\. \d+\.', f'Рисунок {figure_count}.', match)
                        cell.text = cell.text.replace(match, new_text)
                        # Очищаем ячейку и добавляем новый текст
                        # cell.clear()  # Очищаем ячейку
                        # cell.text = cell.text.replace(match, new_text)  # Добавляем новый текст
                        figure_count += 1
            # num_para -= 0
    # print(f'Общее количество рисунков: {figure_count - 1}')
    count = figure_count - 1
    # Сохраняем изменения в документе
    # path = 'откорректированный_' + doc_path
    # rf"temp_files/temp_data/{str(message.chat.id) + '_update_number_picture'}.docx"
    doc.save(doc_path)
    return count
